from __future__ import annotations

import io
from math import ceil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


LABEL_COLUMNS = 3
LABEL_ROWS = 10
LABELS_PER_PAGE = LABEL_COLUMNS * LABEL_ROWS
LABEL_WIDTH = Inches(2.6299)
GUTTER_WIDTH = Inches(0.1201)
LABEL_HEIGHT = Inches(1)
LABEL_FONT_NAME = "Calibri"
LABEL_FONT_SIZE = Pt(12)


def _set_run_font(run) -> None:
    run.bold = True
    run.font.name = LABEL_FONT_NAME
    run.font.size = LABEL_FONT_SIZE
    run._element.rPr.rFonts.set(qn("w:ascii"), LABEL_FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), LABEL_FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), LABEL_FONT_NAME)


def _set_cell_width(cell, width) -> None:
    cell.width = width
    tc_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        cell._tc.get_or_add_tcPr().append(tc_width)
    tc_width.set(qn("w:w"), str(int(width.twips)))
    tc_width.set(qn("w:type"), "dxa")


def _set_table_geometry(table) -> None:
    table.autofit = False
    table_widths = [LABEL_WIDTH, GUTTER_WIDTH, LABEL_WIDTH, GUTTER_WIDTH, LABEL_WIDTH]

    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:w"), str(sum(width.twips for width in table_widths)))
    tbl_width.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_indent = tbl_pr.find(qn("w:tblInd"))
    if tbl_indent is None:
        tbl_indent = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_indent)
    tbl_indent.set(qn("w:w"), "-15")
    tbl_indent.set(qn("w:type"), "dxa")

    cell_margins = tbl_pr.find(qn("w:tblCellMar"))
    if cell_margins is None:
        cell_margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_margins)
    for margin_name in ("left", "right"):
        margin = cell_margins.find(qn(f"w:{margin_name}"))
        if margin is None:
            margin = OxmlElement(f"w:{margin_name}")
            cell_margins.append(margin)
        margin.set(qn("w:w"), "15")
        margin.set(qn("w:type"), "dxa")

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = tbl_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tbl_borders.append(edge)
        edge.set(qn("w:val"), "nil")

    for column, width in zip(table.columns, table_widths):
        column.width = width
    for row in table.rows:
        row.height = LABEL_HEIGHT
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell, width in zip(row.cells, table_widths):
            _set_cell_width(cell, width)


def _fill_label_cell(cell, label_text: str) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    lines = label_text.splitlines()
    if not lines:
        lines = [""]

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1

    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        _set_run_font(run)


def _add_label_page(document: Document, page_labels: list[str]) -> None:
    table = document.add_table(rows=LABEL_ROWS, cols=5)
    _set_table_geometry(table)

    for label_index in range(LABELS_PER_PAGE):
        row_index, label_column = divmod(label_index, LABEL_COLUMNS)
        cell_index = label_column * 2
        label_text = page_labels[label_index] if label_index < len(page_labels) else ""
        _fill_label_cell(table.cell(row_index, cell_index), label_text)

        for gutter_index in (1, 3):
            table.cell(row_index, gutter_index).text = ""


def avery_5520_docx_bytes(labels: list[str]) -> bytes:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.right_margin = Inches(0.1875)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0.1875)

    normal_style = document.styles["Normal"]
    normal_style.font.name = LABEL_FONT_NAME
    normal_style.font.size = LABEL_FONT_SIZE
    normal_style.font.bold = True
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)

    page_count = max(1, ceil(len(labels) / LABELS_PER_PAGE))
    for page_index in range(page_count):
        if page_index:
            document.add_page_break()
        start = page_index * LABELS_PER_PAGE
        _add_label_page(document, labels[start : start + LABELS_PER_PAGE])

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
